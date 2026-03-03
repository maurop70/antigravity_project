import os
import shutil
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
import edge_tts
from pydantic import BaseModel
from typing import Optional

import io
import docx
import pypdf
from coordinator import Coordinator


app = FastAPI(title="Alexa Companion v2")

# CORS Setup
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

coordinator = Coordinator()

# Models
class ChatRequest(BaseModel):
    message: str

class ChatResponse(BaseModel):
    response: str
    active_skill: str

# API Endpoints
@app.post("/api/chat", response_model=ChatResponse)
async def chat_endpoint(request: ChatRequest):
    try:
        print(f"[v2] User: {request.message}")
        response_text = coordinator.process_message(request.message)
        print(f"[v2] Alexa ({coordinator.active_skill_name}): {response_text[:50]}...")
        
        return ChatResponse(
            response=response_text,
            active_skill=coordinator.active_skill_name
        )
    except Exception as e:
        print(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/upload")
async def chat_with_upload(
    message: str = Form(...),
    file: UploadFile = File(...)
):
    try:
        content = ""
        filename = file.filename
        file_bytes = await file.read()
        
        # 1. Determine File Type and Parse
        if filename.lower().endswith(".pdf"):
            try:
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            except Exception as e:
                print(f"PDF Parsing Error: {e}")
                content = f"[Error extracting PDF text: {str(e)}]"
        elif filename.lower().endswith(".docx"):
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    content += para.text + "\n"
            except Exception as e:
                print(f"DOCX Parsing Error: {e}")
                content = f"[Error extracting DOCX text: {str(e)}]"
        else:
            # Try Text Decode
            try:
                content = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                return {
                    "response": f"I received '{filename}', but I couldn't read text from it. I currently support Text, PDF, and Word documents.",
                    "active_skill": coordinator.active_skill_name
                }

        # 2. Augment message
        augmented_message = (
            f"I have uploaded a file named '{filename}'.\n\n"
            f"--- FILE CONTENT START ---\n{content}\n--- FILE CONTENT END ---\n\n"
            f"MY REQUEST:\n{message}"
        )
        
        print(f"[v2] User (File Upload): {filename}")
        
        # 3. Process
        response_text = coordinator.process_message(augmented_message)
        print(f"[v2] Alexa ({coordinator.active_skill_name}): {response_text[:50]}...")
        
        return {
            "response": response_text,
            "active_skill": coordinator.active_skill_name
        }
    except Exception as e:
        print(f"Error in upload: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/upload_knowledge")
async def upload_knowledge(
    file: UploadFile = File(...), 
    category: str = Form(...),
    topic: str = Form(...)
):
    """
    Uploads a file to the Knowledge Base under Category > Topic.
    """
    try:
        content = ""
        filename = file.filename
        file_bytes = await file.read()
        
        # 1. Determine File Type and Parse
        if filename.lower().endswith(".pdf"):
            try:
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            except Exception as e:
                print(f"PDF Parsing Error: {e}")
                content = f"[Error extracting PDF text: {str(e)}]"
        elif filename.lower().endswith(".docx"):
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    content += para.text + "\n"
            except Exception as e:
                print(f"DOCX Parsing Error: {e}")
                content = f"[Error extracting DOCX text: {str(e)}]"
        else:
            # Try Text Decode
            try:
                content = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                raise HTTPException(status_code=400, detail="Cannot read binary file. Please use PDF, Docx, or Text.")

        
        saved_path = coordinator.kb.save_document(category, topic, filename, content)
        
        return {"status": "success", "file": saved_path, "category": category, "topic": topic}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/knowledge_tree")
async def get_knowledge_tree():
    return coordinator.kb.get_hierarchy()

# Parent API Endpoints
class ParentInputRequest(BaseModel):
    content: str
    type: str

@app.get("/api/parent/data")
async def get_parent_data():
    return coordinator.dm.load_data()

@app.post("/api/parent/input")
async def add_parent_input(request: ParentInputRequest):
    try:
        # 1. Analyze with Parent Skill
        parent_skill = coordinator.skills["parent"]
        analysis = parent_skill.analyze_input(request.content, request.type)
        
        # 2. Save to DB
        entry = coordinator.dm.add_parent_input(request.content, request.type, analysis)
        
        # 3. Refresh Context in Coordinator
        coordinator.user_context["parent_data"] = coordinator.dm.load_data()
        
        return entry
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# TTS Endpoint
class SpeakRequest(BaseModel):
    text: str

@app.post("/api/speak")
async def speak(request: SpeakRequest):
    try:
        # Use a friendly, youthful voice: Microsoft Ava (en-US-AvaNeural) - Matches v1
        communicate = edge_tts.Communicate(request.text, "en-US-AvaNeural")
        
        audio_stream = io.BytesIO()
        async for chunk in communicate.stream():
            if chunk["type"] == "audio":
                audio_stream.write(chunk["data"])
        
        audio_stream.seek(0)
        return StreamingResponse(audio_stream, media_type="audio/mpeg")
    except Exception as e:
        print(f"TTS Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Static Files (Frontend)
app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8002)
